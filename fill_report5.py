#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
填写实验五报告：语法制导翻译和中间代码生成
"""

import shutil, os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE    = "/Users/panshuo/first-repo"
SRC     = os.path.join(BASE, "实验五 中间代码生成程序设计.docx")
DEST    = os.path.join(BASE, "实验五_中间代码生成程序设计_完整版.docx")
IMG_DIR = os.path.join(BASE, "_imgs5")

IMG_ARITH       = os.path.join(IMG_DIR, "part1_arith.png")
IMG_BOOL        = os.path.join(IMG_DIR, "part2_bool.png")
IMG_STMT        = os.path.join(IMG_DIR, "part3_stmt.png")
IMG_TOKENIZE    = os.path.join(IMG_DIR, "code_tokenize.png")
IMG_QUAD        = os.path.join(IMG_DIR, "code_quad.png")
IMG_ARITH_CORE  = os.path.join(IMG_DIR, "code_arith_core.png")
IMG_BOOL_CORE   = os.path.join(IMG_DIR, "code_bool_core.png")


# ─────────────── 工具函数 ────────────────────

def del_para(p):
    p._p.getparent().remove(p._p)

def _make_text_elem(cell, text, bold=False):
    """生成段落 XML 元素（先 add 再 detach，以便手动定位插入）"""
    p = cell.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    p._p.getparent().remove(p._p)
    return p._p

def _make_img_elem(cell, img_path, width_in=5.5):
    if not os.path.exists(img_path):
        print(f"  ⚠ 图片不存在: {img_path}")
        return None
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_path, width=Inches(width_in))
    p._p.getparent().remove(p._p)
    return p._p

def insert_after(anchor_p, *elems):
    """在 anchor_p 之后顺序插入各 elem（跳过 None）"""
    cur = anchor_p._p
    for elem in elems:
        if elem is not None:
            cur.addnext(elem)
            cur = elem

def add_text(cell, text, bold=False):
    p = cell.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    return p

def add_img(cell, img_path, width_in=5.5):
    if not os.path.exists(img_path):
        print(f"  ⚠ 图片不存在: {img_path}")
        return
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_path, width=Inches(width_in))
    return p

def find_para(cell, keyword):
    for p in cell.paragraphs:
        if keyword in p.text:
            return p
    return None

def replace_placeholder(cell, keyword, fill_fn):
    target = find_para(cell, keyword)
    if target is None:
        print(f"  ⚠ 未找到占位符: {keyword[:50]!r}")
        return
    before = len(cell.paragraphs)
    fill_fn(cell)
    new_paras = cell.paragraphs[before:]
    parent = target._p.getparent()
    for np in new_paras:
        parent.remove(np._p)
        target._p.addprevious(np._p)
    del_para(target)


# ─────────────── 方法/步骤：5 个问题的答案 ────

ANS_Q1 = """\
答：为简单算术表达式赋值语句设计的产生式（消除左递归后的文法）：

  S   → id = E
  E   → T E'
  E'  → + T E'  |  - T E'  |  ε
  T   → F T'
  T'  → * F T'  |  / F T'  |  ε
  F   → ( E )  |  id  |  num

对应的语义规则（综合/继承属性，.place 记录存放结果的地址）：

  F  → id        F.place = id（标识符名）
  F  → num       F.place = num（数字字面量）
  F  → ( E )     F.place = E.place
  T  → F T'      T'.inh = F.place;  T.place = T'.syn
  T' → * F T'1   t = newtemp();  emit(*, T'.inh, F.place, t)
                 T'1.inh = t;  T'.syn = T'1.syn
  T' → / F T'1   t = newtemp();  emit(/, T'.inh, F.place, t)
                 T'1.inh = t;  T'.syn = T'1.syn
  T' → ε         T'.syn = T'.inh
  E  → T E'      E'.inh = T.place;  E.place = E'.syn
  E' → + T E'1   t = newtemp();  emit(+, E'.inh, T.place, t)
                 E'1.inh = t;  E'.syn = E'1.syn
  E' → - T E'1   t = newtemp();  emit(-, E'.inh, T.place, t)
                 E'1.inh = t;  E'.syn = E'1.syn
  E' → ε         E'.syn = E'.inh
  S  → id = E    emit(=, E.place, 0, id)
"""

ANS_Q2 = """\
答：临时变量由 SymbolTable.new_temp() 统一管理。
  · 使用全局计数器 _temp_cnt，每次调用时自增 1，生成 t1, t2, t3, … 的名称。
  · 用户程序变量在符号表中以正整数下标索引（从 1 开始），
    临时变量以负整数下标（-1, -2, …）区分，两者互不干扰。
  · 每条四元组直接使用变量名（如 t1）作为地址表示；
    如需输出下标形式，可通过符号表的 index() 方法查询。
  · 临时变量不需要显式回收——一个简单的递归下降翻译中，
    临时变量数量等于四元组数量，不存在复用需求。
"""

ANS_Q3 = """\
答：为布尔表达式设计的产生式（优先级：not > and > or，消除左递归后）：

  B   → B1 B'
  B'  → or M B1 B'  |  ε
  B1  → B2 B1'
  B1' → and M B2 B1'  |  ε
  B2  → not B2  |  ( B )  |  E relop E  |  true  |  false
  M   → ε       （语义动作：M.quad = nextquad，标记当前四元组编号）

对应的语义规则（回填属性：truelist / falselist）：

  B2 → E1 relop E2：
      t_idx = emit(j_relop, E1, E2, '?')  ——  truelist = [t_idx]
      f_idx = emit(j, _, _, '?')          ——  falselist = [f_idx]
  B2 → not B2'：    truelist = B2'.falselist;   falselist = B2'.truelist
  B2 → true：       emit(j, _, _, '?');  truelist = [该编号];  falselist = []
  B2 → false：      emit(j, _, _, '?');  truelist = [];  falselist = [该编号]
  B2 → ( B )：      truelist = B.truelist;  falselist = B.falselist

  B1' → and M B2 B1'1（and 短路）：
      backpatch(B2_left.truelist, M.quad)   ——  左操作数为真则计算右操作数
      truelist  = B1'1.syn_tl
      falselist = merge(B2_left.falselist, B1'1.syn_fl)

  B' → or M B1 B'1（or 短路）：
      backpatch(B1_left.falselist, M.quad)  ——  左操作数为假则计算右操作数
      truelist  = merge(B1_left.truelist, B'1.syn_tl)
      falselist = B'1.syn_fl
"""

ANS_Q4 = """\
答：拉链、并链、回填在本实验中的体现：

  · 拉链（链表结构）：
    truelist 和 falselist 均为 Python list，存储所有需要回填同一目标的四元组编号。
    例如 "a<b or c>d"，truelist=[0, 2] 表示四元组 0 和 2 的跳转目标都待填为"真分支入口"。

  · 并链（合并操作）：
    merge(l1, l2) = l1 + l2，将两个待回填列表拼接。
    例如 and 表达式的 falselist = merge(左操作数.falselist, 右操作数.falselist)，
    两者同为假时均跳转到同一目标，合并后一次 backpatch 即可全部填入。

  · 回填（backpatch）：
    QuadList.backpatch(lst, target) 遍历 lst 中每个编号 idx，
    将 _quads[idx] 的第 4 个分量（跳转目标）由 '?' 替换为具体的四元组编号 target。
    时机：当跳转目标（如真分支入口、else 入口）的四元组编号确定后立即执行。

  · M 标记的作用：
    M → ε 是一个"无代价"产生式，其语义动作是捕获当前 nextquad。
    在 and/or 表达式翻译时，翻译右操作数之前先记录 M.quad，
    再用它作为 backpatch 的目标，实现短路跳转的精确定位。
"""

ANS_Q5 = """\
答（选做）：为条件/循环语句设计的产生式：

  Stmts    → Stmt Stmts  |  ε
  Stmt     → if ( B ) { Stmts } ElsePart
           | while ( B ) { Stmts }
           | id = E ;
  ElsePart → else { Stmts }  |  ε

对应的语义规则：

  if-else 语句：
      true_start = nextquad
      backpatch(B.truelist, true_start)   ——  条件为真进入 S1
      翻译 S1，末尾 goto_idx = emit(j, _, _, '?')
      false_start = nextquad
      backpatch(B.falselist, false_start) ——  条件为假进入 S2（或 if 之后）
      翻译 S2（若有 else）
      Stmt.nextlist = merge(S1.nextlist, [goto_idx], S2.nextlist)

  while 语句：
      loop_start = nextquad              ——  记录布尔表达式起始（用于回跳）
      翻译 B，得到 B.truelist, B.falselist
      body_start = nextquad
      backpatch(B.truelist, body_start)  ——  条件为真进入循环体
      翻译循环体 S
      backpatch(S.nextlist, loop_start)  ——  循环体中 break 等回到循环头
      emit(j, _, _, loop_start)          ——  循环体末尾无条件跳回
      Stmt.nextlist = B.falselist        ——  条件为假时退出循环，待回填
"""


# ─────────────── 其余内容文本 ─────────────────

PROC_STEP1 = """\
步骤一：词法分析（tokenize）

整个程序的入口是 tokenize()，将输入字符串扫描为 Token 列表。
实现方式：用正则逐字符匹配，按优先级依次识别数字、标识符/关键字、
双字符关系运算符（<=、>=、==、!=）、单字符运算符及括号。
关键字（and/or/not/true/false/if/else/while）在词法层直接分类，
避免后续解析器为区分关键字和标识符额外处理。
输出 Token 列表末尾追加 EOF 标记，简化解析器的边界判断。
"""

PROC_STEP2 = """\
步骤二：符号表（SymbolTable）与四元组管理（QuadList）

SymbolTable 负责变量/临时变量的统一管理：
· lookup(name)：按需注册用户变量，返回正整数下标（从 1 开始）。
· new_temp()：每次调用全局计数器 _temp_cnt 自增，生成 t1/t2/… 名称，
  返回名称和负整数下标，与用户变量不重叠。

QuadList 管理四元组序列，核心是两个方法：
· emit(op, arg1, arg2, result)：追加一条四元组，返回其编号（0 起）。
  四元组格式为 (op, arg1, arg2, result)，跳转目标暂用 '?' 占位。
· backpatch(lst, target)：遍历编号列表 lst，将对应四元组的第 4 项
  由 '?' 替换为具体目标编号 target，实现回填。
"""

PROC_STEP3 = """\
步骤三：算术表达式翻译（ArithParser）

采用递归下降 + 综合/继承属性翻译，核心是 parse_Eprime(inh) 和 parse_Tprime(inh)：
· inh 为继承属性，传递左操作数的地址（变量名或临时变量名）。
· 每识别一个运算符，调用 new_temp() 申请新临时变量，
  立即 emit 一条四元组，并将新临时变量作为新的 inh 向右递归。
· parse_S() 最后 emit(=, E.place, 0, id) 完成赋值。

调试过程：先测试 a=b+c（只有加法），验证 emit 顺序正确；
再加入 a=b*c+d，确认乘法优先级体现在先生成 (*,b,c,t1) 再生成 (+,t1,d,t2)；
最后加括号用例 result=(a+b)*(c-d) 验证括号改变优先级的效果。
"""

PROC_STEP4 = """\
步骤四：布尔表达式翻译（BoolParser，回填法）

将 B 拆分为三层（体现优先级 not > and > or）：
· parse_B / parse_Bprime：处理 or，回填左操作数的 falselist 到右操作数入口。
· parse_B1 / parse_B1prime：处理 and，回填左操作数的 truelist 到右操作数入口。
· parse_B2：处理 not、括号、关系表达式和 true/false。

关键实现细节：
· M 标记不单独写解析函数，而是在调用右操作数之前直接用
  m_quad = quad.next_quad() 捕获当前编号，作为 backpatch 目标。
· parse_B2 中识别关系表达式时，临时构造 ArithParser 共用 tokens，
  解析完算术子式后将 arith.pos 同步回 self.pos，避免 pos 错位。

调试过程：手工追踪 a<b and c>d 的四元组生成顺序，
逐步验证 truelist/falselist 的变化，确认回填目标正确。
"""

CONCLUSION_ARITH = """\
测试用例及四元组输出：

(1) 输入: a=b+c*e/g
    (0) (*, c, e, t1)
    (1) (/, t1, g, t2)
    (2) (+, b, t2, t3)
    (3) (=, t3, 0, a)

(2) 输入: result=(a+b)*(c-d)
    (0) (+, a, b, t1)
    (1) (-, c, d, t2)
    (2) (*, t1, t2, t3)
    (3) (=, t3, 0, result)

(3) 输入: z=((a+b)*c-d)/e+f
    (0) (+, a, b, t1)
    (1) (*, t1, c, t2)
    (2) (-, t2, d, t3)
    (3) (/, t3, e, t4)
    (4) (+, t4, f, t5)
    (5) (=, t5, 0, z)

验证：按运算符优先级（先乘除后加减，括号最高），临时变量编号从 t1 顺序增加，输出正确。
完整测试截图见下图：
"""

CONCLUSION_BOOL = """\
测试用例：

(1) a < b and c > d
    (0) (j<, a, b, 2)     ← a<b 为真跳到 (2)，即 and 右侧
    (1) (j, _, _, ?)      ← a<b 为假，falselist=[1]，待回填
    (2) (j>, c, d, ?)     ← c>d 为真，truelist=[2]，待回填
    (3) (j, _, _, ?)      ← c>d 为假，falselist=[1,3]

(2) a < b or c > d
    (0) (j<, a, b, ?)     ← a<b 为真，truelist=[0,2]
    (1) (j, _, _, 2)      ← a<b 为假跳到 (2)，即 or 右侧
    (2) (j>, c, d, ?)
    (3) (j, _, _, ?)      ← falselist=[3]

(3) not (a<b or c>d) and e<=f
    truelist=[4], falselist=[0,2,5]，
    体现了 not 对 truelist/falselist 的交换，以及 and 的拼链操作。

完整测试截图见下图：
"""

CONCLUSION_STMT = """\
(1) if (a < b) { x = a+1; } else { x = b+2; }
    (0) (j<, a, b, 2)     ← B.truelist backpatch → 真分支入口
    (1) (j, _, _, 5)      ← B.falselist backpatch → else 入口
    (2) (+, a, 1, t1)
    (3) (=, t1, 0, x)
    (4) (j, _, _, ?)      ← 真分支末尾 goto，nextlist=[4]
    (5) (+, b, 2, t2)
    (6) (=, t2, 0, x)

(2) while (i < n) { i = i+1; }
    (0) (j<, i, n, 2)     ← 条件真跳循环体
    (1) (j, _, _, ?)      ← falselist=[1]，循环出口
    (2) (+, i, 1, t1)
    (3) (=, t1, 0, i)
    (4) (j, _, _, 0)      ← 循环体末尾回跳 loop_start=0

复合条件 if (a<b and c>0) 截图见下图：
"""

INSIGHT_TEXT = """\
通过本次实验，我对语法制导翻译和中间代码生成有了系统性的理解。

一、语法制导翻译的本质

语法制导翻译将语义动作嵌入文法产生式中，在语法分析的同时计算语义属性。
本实验使用了两类属性：
· 综合属性（S-属性）：如 E.place，由子节点计算，自底向上传递，
  适用于表达式中间代码生成，递归下降中以返回值体现。
· 继承属性（I-属性）：如 E'.inh，由父节点或兄弟节点传递，
  消除左递归后以函数参数形式向下传递（如 parse_Eprime(inh)）。

二、四元组（三地址码）的优点

四元组 (op, arg1, arg2, result) 是一种简洁的中间表示，
临时变量使无限寄存器假设成立，每条指令恰好执行一个操作，
便于后续的代码优化（如公共子表达式消除、常量折叠）和目标代码生成。
相比直接生成汇编，四元组与具体机器无关，抽象层次合适。

三、回填法的精妙之处

布尔表达式的翻译面临"跳转目标未知"的问题——翻译 B 时还不知道
真/假分支从哪条指令开始，因此先生成含 ? 的跳转四元组，
用 truelist/falselist 记录待填入位置，等目标确定后再 backpatch。
这一延迟求值的思想优雅地解决了单遍翻译中的前向引用问题。
M 标记（M→ε）在文法中插入语义动作点、捕获 nextquad，
让回填目标的传递完全融入文法规则，无需额外的数据结构。

四、条件/循环语句的翻译规律

if-else 和 while 的翻译实质是对 B.truelist/falselist 的定向回填：
· truelist → 真分支入口
· falselist → 假分支入口（if 的 else 或 while 的出口）
while 在循环体末尾加 goto loop_start，实现循环，S.nextlist 向外传递。
这一模式可以递归嵌套，体现了语法制导翻译处理控制流的通用性。

五、实现体会

递归下降分析天然契合综合/继承属性的传递，无需显式建树，是小型编译器的实用选择。
回填法的调试需要仔细追踪每条四元组的编号和 truelist/falselist 的演变，
建议手工模拟几个简单用例后再运行程序验证，可以快速定位拼链和回填方向错误。
"""


# ─────────────── 主逻辑 ──────────────────────

def main():
    print("▶ 复制模板 …")
    shutil.copy2(SRC, DEST)
    doc = Document(DEST)
    tbl = doc.tables[0]

    row1 = tbl.rows[1].cells[0]   # 方法/步骤
    row2 = tbl.rows[2].cells[0]   # 实验过程
    row3 = tbl.rows[3].cells[0]   # 实验结论
    row4 = tbl.rows[4].cells[0]   # 心得体会

    # ── 行[1] 方法/步骤：每个问题后插入答案 ───
    Q_ANS = [
        ("你为简单算术表达式、赋值语句设计的产生式和语义规则分别是什么", ANS_Q1),
        ("翻译输出过程中，需要生成多个临时变量，你是如何管理、控制这些临时变量的", ANS_Q2),
        ("针对选做部分的布尔表达式，你为布尔表达式设计的产生式和相应的语义规则分别是什么", ANS_Q3),
        ("布尔表达式翻译过程中，出现的拉链、并链、回填等操作", ANS_Q4),
        ("选做）针对选做部分的条件语句，你为条件语句设计的产生式和相应的语义规则分别是什么", ANS_Q5),
    ]
    for keyword, answer in Q_ANS:
        qp = find_para(row1, keyword)
        if qp is None:
            print(f"  ⚠ 未找到问题段落: {keyword[:40]!r}")
            continue
        ans_elem = _make_text_elem(row1, answer)
        insert_after(qp, ans_elem)
    print("  ✓ 方法/步骤")

    # ── 行[2] 实验过程 ───────────────────────
    # 替换 P1 占位符，并额外删除 P2（"例如，为了编程调试方便…"）
    def fill_proc(cell):
        add_text(cell, PROC_STEP1)
        add_img(cell, IMG_TOKENIZE, width_in=5.8)
        add_text(cell, PROC_STEP2)
        add_img(cell, IMG_QUAD, width_in=5.8)
        add_text(cell, PROC_STEP3)
        add_img(cell, IMG_ARITH_CORE, width_in=5.8)
        add_text(cell, PROC_STEP4)
        add_img(cell, IMG_BOOL_CORE, width_in=5.8)

    replace_placeholder(row2, "实验过程及内容，除了代码设计说明", fill_proc)
    # 删除剩余的说明性占位段落
    for p in list(row2.paragraphs):
        if "为了编程调试方便" in p.text or "请在这里，分步骤" in p.text:
            del_para(p)
    print("  ✓ 实验过程")

    # ── 行[3] 实验结论 ───────────────────────
    # 三个小节各自：在节标题后插入内容，删除"请…"占位段
    sec_map = [
        ("简单算术表达式的翻译的实验结果",
         "请设计测试用例，或者选用、截OJ上通过测试的截图进行说明",
         CONCLUSION_ARITH, IMG_ARITH),
        ("布尔表达式的翻译的实验结果",
         "请设计测试用例，或者选用、请截OJ上通过测试的截图进行说明",
         CONCLUSION_BOOL, IMG_BOOL),
        ("（选做）条件语句的翻译的实验结果",
         "请自行设计恰当的输入，并检验输出结果",
         CONCLUSION_STMT, IMG_STMT),
    ]
    for header_kw, placeholder_kw, text, img in sec_map:
        header_p = find_para(row3, header_kw)
        placeholder_p = find_para(row3, placeholder_kw)
        if header_p is None:
            print(f"  ⚠ 未找到节标题: {header_kw[:40]!r}")
            continue
        # 在节标题后插入答案
        text_elem = _make_text_elem(row3, text)
        img_elem  = _make_img_elem(row3, img, width_in=5.8)
        insert_after(header_p, text_elem, img_elem)
        # 删除占位段
        if placeholder_p is not None:
            del_para(placeholder_p)
    print("  ✓ 实验结论")

    # ── 行[4] 心得体会 ───────────────────────
    def fill_insight(cell):
        add_text(cell, INSIGHT_TEXT)

    replace_placeholder(row4, "除了实验的心得体会外", fill_insight)
    print("  ✓ 心得体会")

    doc.save(DEST)
    print(f"\n✅ 完成！输出：{DEST}")


def _make_text_elem(cell, text, bold=False):
    p = cell.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10.5)
    p._p.getparent().remove(p._p)
    return p._p

def _make_img_elem(cell, img_path, width_in=5.5):
    if not os.path.exists(img_path):
        print(f"  ⚠ 图片不存在: {img_path}")
        return None
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img_path, width=Inches(width_in))
    p._p.getparent().remove(p._p)
    return p._p


if __name__ == '__main__':
    main()
