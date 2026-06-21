#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
填写计算机伦理期末大作业答题纸
"""

import os, shutil, zipfile, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

BASE = "/Users/panshuo/first-repo"
TEMPLATE = os.path.join(BASE, "深圳大学2026《计算机伦理》期末大作业-姓名-学号--考核专用答题纸 - 模板.dotx")
DEST = os.path.join(BASE, "深圳大学2026《计算机伦理》期末大作业_医疗AI伦理分析_完整版.docx")
TMP_DOCX = "/tmp/ethics_template.docx"


# ─── 把 .dotx 的 Content-Type 改为 .docx 再打开 ────────────
def dotx_to_docx(src, dst):
    with zipfile.ZipFile(src, 'r') as zin:
        names = zin.namelist()
        with zipfile.ZipFile(dst, 'w', zipfile.ZIP_DEFLATED) as zout:
            for name in names:
                data = zin.read(name)
                if name == '[Content_Types].xml':
                    data = data.replace(
                        b'wordprocessingml.template.main+xml',
                        b'wordprocessingml.document.main+xml'
                    )
                zout.writestr(name, data)

dotx_to_docx(TEMPLATE, TMP_DOCX)
doc = Document(TMP_DOCX)


# ─── 格式工具 ──────────────────────────────────────────────
def set_run_fmt(run, size_pt=9, bold=False, font_name='宋体'):
    run.font.size = Pt(size_pt)
    run.font.name = font_name
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def set_para_fmt(para, space_before=0, space_after=0,
                 first_line_indent=None, align=None):
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    pf.line_spacing = Pt(12)          # 单倍行距（五号字 9pt × ~1.3）
    if first_line_indent is not None:
        pf.first_line_indent = Pt(first_line_indent)
    if align is not None:
        para.alignment = align

def add_para(cell_or_doc, text, bold=False, size=9,
             indent=None, align=None, heading=False):
    """向 cell 或 doc 追加一段，返回 paragraph"""
    if hasattr(cell_or_doc, 'add_paragraph'):
        p = cell_or_doc.add_paragraph()
    else:
        p = cell_or_doc.add_paragraph()

    run = p.add_run(text)
    set_run_fmt(run, size_pt=size, bold=bold)
    set_para_fmt(p, space_before=0, space_after=0,
                 first_line_indent=indent, align=align)
    return p

def add_heading(doc, text, level=1):
    """用正文段落模拟标题（粗体 + 适当缩进）"""
    size = {1: 10.5, 2: 9, 3: 9}.get(level, 9)
    p = add_para(doc, text, bold=True, size=size,
                 indent=0 if level == 1 else 14)
    return p


# ─── 找到答题区起始位置并清除占位符 ──────────────────────
# 模板段落结构：…题目：… (开始答卷。。。)
answer_start_idx = None
title_para_idx = None

for i, p in enumerate(doc.paragraphs):
    if '题目：' in p.text:
        title_para_idx = i
    if '开始答卷' in p.text:
        answer_start_idx = i
        break

# 填写题目
if title_para_idx is not None:
    tp = doc.paragraphs[title_para_idx]
    for run in tp.runs:
        run.clear()
    run = tp.add_run(
        "题目：医疗场景中AI智能检测与疾病预测的计算机伦理问题分析"
    )
    set_run_fmt(run, size_pt=9, bold=False)

# 删除 (开始答卷。。。) 占位段
if answer_start_idx is not None:
    p = doc.paragraphs[answer_start_idx]
    p._element.getparent().remove(p._element)


# ─── 正文内容 ──────────────────────────────────────────────
d = doc   # 直接向 doc 追加段落（答题区在文档末尾）

SECTIONS = [

("引言",
"""人工智能技术在医疗健康领域的广泛应用，标志着一场深刻的技术与社会变革正在发生。从医学影像辅助诊断到疾病风险预测，从智能分诊系统到个性化治疗建议，AI系统正在重塑传统医疗服务的形态与边界。然而，医疗场景的特殊性决定了其中的伦理问题远比一般技术应用更为复杂：医疗数据的高度敏感性、患者生命健康的不可替代价值，以及AI决策对就医行为与治疗方案的潜在影响，使得医疗AI不仅是一个技术问题，更是一个关乎公平、隐私、责任与人类尊严的伦理命题。

本报告以计算机伦理为核心框架，结合课程讲义内容与真实案例，对医疗AI的社会影响、监管平衡、伦理问题及治理策略进行系统分析，并提出个人主张与建议。"""),

("一、医疗AI智能检测与疾病预测的社会影响分析（20分）",
None),

("（一）正面影响",
None),

("1. 提高诊断效率与准确率",
"""AI辅助诊断系统能够在短时间内处理海量医学影像数据，大幅提升诊断效率。Google Health开发的基于深度学习的眼底病变检测系统，在糖尿病视网膜病变筛查中达到了超过90%的灵敏度，与专科医生水平相当，并已在印度、泰国等医疗资源匮乏地区的实际临床场景中得到验证（Gulshan et al., JAMA, 2016）。AI系统不受疲劳影响，在24小时高吞吐量筛查场景中具有明显优势，可有效减少人为漏诊。"""),

("2. 促进医疗资源的普惠化与公平获取",
"""优质医疗资源的地理分布不均是全球性问题。AI辅助诊断系统可以部署在基层医疗机构，使边远地区患者无需长途就医即可获得接近三甲医院水平的初步诊断。腾讯觅影、阿里健康等国内AI医疗平台已将部分功能延伸至县域医院，为基层医生提供AI辅助诊断支持，在一定程度上缓解了"看病难"问题。"""),

("3. 推动慢性病管理与疾病预防",
"""疾病风险预测模型可以整合患者的电子病历、实验室检验指标、生活方式数据与可穿戴设备数据，提前识别高风险人群。心脑血管疾病、糖尿病、癌症等慢性病的早期预警，能够促使患者及时干预、改变生活方式，从而显著降低疾病进展风险。美国Epic系统内嵌的败血症早期预警模型，已在ICU场景中显著提前了危重症的识别时间。"""),

("4. 辅助临床决策，减少医疗差错",
"""临床决策支持系统（CDSS）能够在医生开具处方时，实时提示药物相互作用、禁忌症及剂量异常，减少由疲劳、信息不足或认知偏差导致的医疗差错。研究显示，AI辅助的放射科报告系统可将漏诊率降低约20%（McKinney et al., Nature, 2020）。"""),

("5. 加速新药研发",
"""AI技术在药物靶点发现、分子结构预测（如DeepMind AlphaFold2）以及临床试验患者招募等方面展现出巨大潜力，显著缩短了从基础研究到临床应用的周期，有望降低新药研发成本，最终使患者获益。"""),

("（二）负面影响",
None),

("1. 算法歧视与医疗不公平",
"""训练数据的偏差是医疗AI最严重的系统性风险之一。2019年，Obermeyer等研究者在《科学》发表研究，揭示Optum公司一款广泛使用的医疗风险分层算法存在严重种族偏见：该算法以"历史医疗花费"作为健康需求的代理变量，而由于历史上黑人患者获得医疗服务的机会较少，其花费往往低于同等健康状况的白人患者，导致算法系统性地低估了黑人患者的医疗需求，使黑人患者被推荐接受额外健康管理项目的比例下降约50%。"""),

("2. 患者隐私泄露与数据安全风险",
"""医疗AI系统需要大规模收集和处理患者的敏感个人健康数据。2017年，英国NHS皇家自由医院在未获得患者充分知情同意的情况下，将约160万患者的个人健康数据提供给DeepMind开发Streams急性肾损伤预警应用。英国ICO认定该数据共享行为违反了《数据保护法》，此案成为医疗AI数据隐私争议的标志性案例。"""),

("3. 医疗决策责任归属模糊",
"""当AI辅助诊断给出错误建议并导致不良医疗结果时，责任主体难以明确：是AI开发者、医疗机构、部署系统的医生，还是患者本身？这种责任归属的模糊性可能导致患者权益受损时无法获得有效救济，也可能使医生过度依赖AI系统而丧失独立判断能力（"自动化偏见"问题）。"""),

("4. 加剧数字健康鸿沟",
"""AI医疗服务往往依赖智能手机、可穿戴设备和高速网络，而老年人、农村居民、低收入群体等往往缺乏这些技术条件。如果医疗AI的红利主要由受过良好教育、具备技术能力的城市中产阶级享有，则可能进一步扩大不同社会群体之间的健康不平等。"""),

("5. 对医患关系与医生职业角色的冲击",
"""当患者接受到AI生成的疾病预测结果（如"您患某种疾病的风险为75%"）时，可能产生不必要的心理焦虑或过度医疗，而此类情感与心理层面的需求是AI系统目前难以有效响应的。医疗AI的广泛应用若处理不当，可能弱化基于信任与沟通建立的医患关系。"""),

("二、医疗AI发展中技术创新、临床安全与政策监管的平衡（15分）",
None),

("（一）加强监管的必要性与潜在风险（8分）",
None),

("1. 加强监管的必要性",
"""医疗AI直接涉及患者生命健康，错误的诊断结果可能延误最佳治疗时机，造成不可逆伤害。美国FDA在其2021年发布的《AI/ML医疗器械软件行动计划》中明确提出了对AI医疗软件全生命周期监管要求，包括算法透明度、性能持续监测与变更报告。欧盟《人工智能法案》（2024年正式通过）将用于医疗目的的高风险AI系统列为重点监管对象，要求满足高质量训练数据、透明度与可解释性、人类监督及上市后监测等要求。中国NMPA也在人工智能医疗器械注册审评中要求提供详细的算法性能验证报告与临床试验数据。这些监管框架的建立，是对医疗AI商业化扩张中可能出现的风险进行系统性防范的必要举措。"""),

("2. 过度监管的潜在风险",
"""繁琐的审批程序可能显著延长AI医疗产品的上市周期，使本可造福患者的创新技术迟迟无法进入临床应用。高昂的合规成本对大型企业可能尚可接受，但对初创公司和学术研究机构则可能形成进入壁垒，抑制多元化的创新生态。此外，若监管规则过于僵化，难以跟上AI技术的快速迭代节奏，将导致"监管滞后"问题，使实际应用中的新型风险无法得到有效约束。"""),

("（二）鼓励技术创新与开放应用的益处及可能问题（7分）",
None),

("1. 开放创新的益处",
"""适度宽松的政策环境有助于吸引资本和人才投入医疗AI领域，促进算法性能的持续提升与应用场景的快速拓展。COVID-19疫情期间，多家AI公司迅速开发出基于胸部CT的新冠肺炎辅助诊断系统，在核酸检测能力有限的早期阶段辅助了大量疑似病例的快速分诊，展示了AI快速响应的潜力。开放的数据政策（如允许医疗机构在脱敏条件下共享数据用于科研）有助于构建更大规模、更具代表性的训练数据集，从根本上缓解数据偏见问题。"""),

("2. 可能带来的问题",
"""缺乏足够验证的AI系统一旦进入临床应用，可能产生严重后果。IBM Watson for Oncology的肿瘤治疗建议被发现与医学专家意见存在重大分歧，部分建议被指出"不安全"（Strickland, IEEE Spectrum, 2019）。过于宽松的数据共享政策可能使商业利益凌驾于患者隐私保护之上，竞争驱动的"数据竞赛"也可能加剧大型科技公司对医疗数据的垄断。

因此，理想的政策框架应当是"有边界的开放"：在保障核心安全底线与数据权益的前提下，为技术创新预留足够空间，并建立适应技术发展的动态监管机制。"""),

("三、从计算机伦理角度分析关键伦理问题（30分）",
None),

("（一）算法偏见与公平性（10分）",
None),

("1. 成因分析",
"""课程讲义"算法伦理与社会公平"一章指出，公平性是计算机伦理的核心价值之一，涵盖机会平等与结果平等两个层面。对于医疗AI而言，算法偏见的根源通常存在于三个环节：

第一，训练数据的历史偏差。若训练数据反映了历史上医疗资源分配的不公平现实，模型将习得并固化这种偏差。Obermeyer等（2019）的研究揭示，以"历史医疗花费"作为健康需求代理变量，会将社会不平等"编码"进算法逻辑之中。

第二，代表性不足（underrepresentation）。大多数医学影像AI模型主要基于来自发达国家和特定人群的数据进行训练，在应用于不同肤色、体型或疾病表现模式的人群时，性能可能显著下降。Adamson & Smith（2018）在《JAMA Dermatology》的研究表明，某些皮肤病AI诊断系统在深色皮肤患者中的准确率明显低于浅色皮肤患者。

第三，代理变量的选择偏差。选择与健康需求高度相关但内嵌社会偏见的代理变量作为训练标签，会将不平等系统性地引入模型预测。"""),

("2. 缓解策略",
"""在数据收集阶段，应主动进行多元化数据采集，确保不同人口统计学群体的充分代表；在模型设计阶段，采用去偏（debiasing）算法和公平性约束，引入多维度公平性评估指标（如equalized odds、demographic parity）；在部署阶段，建立针对不同人群的分组性能报告制度（disaggregated performance reporting），持续监测潜在偏见的演化。"""),

("（二）AI职业伦理与工程师角色（10分）",
None),

("1. 工程师的职业伦理困境",
"""课程讲义"IT职业道德与社会责任"一章强调，IT专业人员不仅承担技术实现的职责，更是社会资源分配与价值判断的实际参与者。在医疗AI开发中，工程师面临效率与公平的深层张力：若以"最大化整体预测准确率"为优化目标，模型在多数群体上表现优异，但可能系统性地牺牲少数群体的预测精度——从纯粹技术指标角度看是"最优"的，从社会公平角度看则是有偏见的。

工程师还面临商业压力与伦理责任之间的冲突：当雇主要求加快产品上市节奏，而工程师认为模型尚未经过足够验证时，是否有责任发声？ACM《软件工程师职业道德规范》（2018）原则1.01明确指出：工程师有责任将公众健康与安全置于雇主和客户利益之上。在医疗AI领域，这意味着工程师不能以"只是工具实现者"为由规避对算法社会影响的责任追究。"""),

("2. 伦理判断力作为专业成熟度的核心",
"""成熟的技术专业人员不仅能够构建高性能模型，还能够识别模型可能造成的非预期后果，并在技术设计阶段就将伦理考量嵌入决策流程。这要求工程师具备跨学科素养，能够与伦理学家、医学专家、患者代表及政策制定者进行有效协作，形成"价值敏感设计"（value-sensitive design）的实践范式。从"工具实现者"到"社会资源分配的设计者"的角色转变，是IT行业职业成熟度提升的核心标志。"""),

("（三）全生命周期治理与多层次责任结构（10分）",
None),

("1. 四个层面的风险分析",
"""课程讲义"信息技术治理与责任"一章指出，复杂技术系统的伦理风险往往不是静态的，而是随着技术的开发、部署与使用动态演化的。以"联邦学习联合诊断模型"为例，该技术通过在多个医疗机构间协作训练模型而不集中传输原始数据，其风险可能出现在以下四个层面：

代码/模型层面：联邦学习中各参与节点的本地模型更新可能引入数据投毒攻击（data poisoning），恶意参与者可以通过操纵本地训练数据影响全局模型的学习方向；模型本身的可解释性不足，使临床医生难以理解特定预测结果的依据。

系统/产品层面：联邦学习系统的聚合算法设计、安全协议实现和差分隐私参数设置，决定了整体系统对隐私泄露和对抗性攻击的抵御能力。产品部署中的配置错误或版本管理不善，可能引入额外的系统性风险。

组织层面：参与联邦学习的各医疗机构在数据质量控制、合规管理和内部治理方面存在差异，可能导致"木桶效应"——整体系统的安全与公平性受制于最薄弱的参与节点。机构间对数据使用目的与商业利益的理解差异，也可能造成责任边界模糊。

社会层面：联邦学习模型一旦规模化应用，其对特定疾病诊断路径的影响可能产生系统性的医疗资源再分配效应；若不同地区的医疗机构参与比例差异悬殊，可能进一步巩固既有的医疗不平等格局。"""),

("2. 通过工程手段实现责任界定与追溯",
"""为实现有效的责任界定，可采取以下工程手段：建立可审计的模型训练日志（audit trail），记录每次迭代中各参与节点的贡献与模型状态；引入联邦学习的可追溯性框架，确保模型更新的来源可被追溯；制定明确的多方数据合作协议，在法律与技术双重层面划定责任边界；并建立独立的第三方审计机制，对模型性能与偏见进行定期评估。"""),

("四、个人观点与案例分析（25分）",
None),

('（一）个人主张：构建“技术边界清晰、责任主体明确、患者中心”的医疗AI伦理治理框架（15分）',
None),

("1. 人类监督不可缺位，AI辅助而非替代",
"""医疗AI的本质定位应当是增强医生的判断能力，而非取代医生的临床决策权。在直接影响患者治疗方案的场景中，应建立强制性的"人类审核节点"（human-in-the-loop），确保最终决策由具有完整信息和临床责任的医生作出。这不仅是对患者安全的保障，也是防止"自动化偏见"侵蚀临床决策质量的必要机制。当AI系统的建议与医生临床判断发生冲突时，医生应有权力也有责任作出最终决定，而不是被迫遵从AI输出。"""),

("2. 透明度与可解释性作为准入门槛",
"""医疗AI系统应当以可解释性（explainability）为基本要求，不仅向技术专家，也向医生和患者提供可理解的决策依据。对于高风险应用（如肿瘤辅助诊断、重症预警），"黑箱"模型即便准确率再高，也不应在缺乏充分解释机制的条件下进入临床应用。可解释AI（XAI）技术的发展——如SHAP值分析、LIME方法、注意力可视化等——为实现临床可解释性提供了技术基础，应将其纳入医疗AI产品的强制性技术规范。"""),

("3. 患者数据权属的重新界定与数据信托模式",
"""医疗数据是患者生命历程的数字化映射，其权属不应默认归属于收集数据的医疗机构或技术公司。应确立患者对其个人医疗数据的基本权利，包括知情权、访问权、更正权和可携带权（data portability）。在此基础上，探索基于患者明确授权的数据信托（data trust）模式，在保障数据合理流通以促进AI研发的同时，确保患者权益不受侵害。参考欧盟GDPR及中国《个人信息保护法》的相关框架，建立医疗数据使用的分层授权与用途限制机制。"""),

("4. 多主体协同治理：平台、医疗机构、医生与患者的责任划分",
"""医疗AI的伦理治理不能依靠单一主体完成。AI平台（开发者）应对算法性能、数据质量与偏见风险承担首要技术责任；医疗机构（部署者）应承担对AI系统临床适用性的验证责任，建立院内AI应用的审查和监督机制；医生（使用者）应接受充分的AI素养培训，理解AI系统的能力边界，避免过度依赖；患者应有权知晓在其诊疗中使用了哪些AI辅助系统，并有权拒绝；监管机构则应提供清晰的准入标准、上市后监测要求和违规处罚机制，形成闭环。"""),

("（二）案例分析（10分）",
None),

("案例一：IBM Watson for Oncology——技术宣传与临床现实的落差",
"""IBM Watson for Oncology被作为"认知计算革命医疗"的旗舰案例在全球多个国家推广，宣称能为肿瘤科医生提供基于循证医学的个性化治疗建议。然而，STAT News的调查报道（Strickland, 2019）与多家合作医院（包括印度Manipal医院）的内部评估显示，Watson的治疗建议与医学专家共识存在严重分歧，部分建议被临床医生认定为不安全。深入调查发现，Watson的训练主要依赖来自纪念斯隆-凯特琳癌症中心的有限病例，缺乏对不同人群、不同医疗体系的充分覆盖，导致在全球化推广中出现严重的泛化失败。

这一案例深刻揭示了医疗AI商业化过程中的伦理风险：过度营销与不充分验证的结合，使尚未成熟的技术进入临床，不仅可能给患者带来实际伤害，还严重损害了公众对医疗AI整体的信任。这说明，监管机构对医疗AI临床有效性的独立验证要求，以及行业自律对夸大宣传的约束，是不可缺少的伦理守护机制。"""),

("案例二：DeepMind Streams与NHS数据共享争议——知情同意的边界",
"""2015年至2017年间，英国Royal Free Hospital（NHS信托机构）与DeepMind签署协议，将约160万NHS患者的个人健康数据（包括病历、实验室结果及部分HIV检测结果等高度敏感信息）提供给DeepMind，用于开发急性肾损伤（AKI）早期预警应用Streams。2017年，英国ICO调查认定：该数据共享协议未获得患者明确知情同意，且数据使用范围超出了开发AKI预警这一特定目的，违反了《数据保护法》的相关规定。此后，谷歌收购DeepMind并将Streams相关资产整合至Google Health，进一步引发了公众对商业科技公司获取NHS敏感数据动机的质疑。

这一案例说明：在医疗AI数据收集与使用过程中，目的限制原则（purpose limitation）与知情同意（informed consent）不是可以为了技术进步而妥协的形式程序，而是保护患者基本权益、维护社会对医疗数字化转型信任的实质性伦理基础。无论技术成果多么有益，绕过患者知情权的数据收集行为都构成对患者自主权的侵犯，平台方、医疗机构和监管机构各有不可推卸的责任。"""),

("案例三：COVID-19中的AI辅助CT诊断——应急响应的伦理张力",
"""2020年新冠疫情暴发初期，多家中国AI公司（如阿里巴巴、科大讯飞等）迅速开发出基于CT影像的新冠肺炎辅助诊断系统，在核酸检测能力有限的早期阶段辅助了大批疑似病例的快速分诊，展示了AI快速响应的潜力。然而，这些系统大多在缺乏充分临床验证的条件下被快速部署，其泛化性能尚不明确。事后独立评估研究发现，部分系统在外部验证集上的性能显著低于开发者自报指标（Roberts et al., Nature Machine Intelligence, 2021）。

这一案例揭示了应急情境下技术创新与临床安全之间的深层张力：在极端公共卫生危机中，适度放宽准入门槛可能是合理的权宜之计，但不应成为长期降低验证标准的先例。应急部署后的全面评估与知识总结，应当被视为配套义务，以避免形成"紧急状态规范化"的不良先例。"""),

("结语",
"""医疗AI的伦理治理，本质上是技术社会中人类如何在追求效率与进步的同时，坚守公平、隐私、安全与人类尊严的价值承诺。当前，没有任何单一主体——无论是开发者、医疗机构、监管机构还是患者——能够独自承担所有责任。唯有构建覆盖技术研发、临床部署、政策监管与公众参与的多层次协同治理框架，才能使医疗AI真正成为造福全体人类的公共卫生资产，而非少数人牟利或技术权力扩张的工具。"""),

("参考文献",
"""[1] Gulshan, V., et al. (2016). Development and Validation of a Deep Learning Algorithm for Detection of Diabetic Retinopathy in Retinal Fundus Photographs. JAMA, 316(22), 2402–2410.
[2] Obermeyer, Z., et al. (2019). Dissecting racial bias in an algorithm used to manage the health of populations. Science, 366(6464), 447–453.
[3] McKinney, S.M., et al. (2020). International evaluation of an AI system for breast cancer screening. Nature, 577, 89–94.
[4] Adamson, A.S., & Smith, A. (2018). Machine Learning and Health Care Disparities in Dermatology. JAMA Dermatology, 154(11), 1247–1248.
[5] Strickland, E. (2019). How IBM Watson Overpromised and Underdelivered on AI Health Care. IEEE Spectrum.
[6] UK Information Commissioner's Office (ICO). (2017). Royal Free – Google DeepMind trial failed to comply with data protection law.
[7] Roberts, M., et al. (2021). Common pitfalls and recommendations for using machine learning to detect and prognosticate for COVID-19 using chest radiographs and CT scans. Nature Machine Intelligence, 3, 199–217.
[8] European Parliament. (2024). EU Artificial Intelligence Act.
[9] U.S. FDA. (2021). Artificial Intelligence/Machine Learning (AI/ML)-Based Software as a Medical Device (SaMD) Action Plan.
[10] ACM. (2018). Software Engineering Code of Ethics and Professional Practice.
[11] 课程讲义：计算机伦理，深圳大学，2025–2026学年第二学期。"""),
]


def is_section_header(text):
    return (text.startswith(('一、','二、','三、','四、','（一）','（二）','（三）',
                              '1.','2.','3.','4.','5.','案例','引言','结语','参考文献'))
            or text.endswith(('分）', '分)')))

for title, body in SECTIONS:
    # 判断标题层级
    lv1 = title.startswith(('一、','二、','三、','四、','引言','结语','参考文献'))
    lv2 = title.startswith(('（一）','（二）','（三）'))
    lv3 = (title[0].isdigit() and title[1] == '.') or title.startswith('案例')

    indent_pts = 0 if lv1 else (14 if lv2 else 18)
    title_bold = True
    title_size = 10.5 if lv1 else (9 if lv2 else 9)

    p_title = d.add_paragraph()
    run_t = p_title.add_run(title)
    set_run_fmt(run_t, size_pt=title_size, bold=title_bold)
    set_para_fmt(p_title, space_before=3 if lv1 else 1,
                 space_after=0, first_line_indent=indent_pts)

    if body:
        for para_text in body.split('\n\n'):
            para_text = para_text.strip()
            if not para_text:
                continue
            # 段内换行保留
            for line in para_text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                p_body = d.add_paragraph()
                run_b = p_body.add_run(line)
                set_run_fmt(run_b, size_pt=9)
                set_para_fmt(p_body, space_before=0, space_after=0,
                             first_line_indent=18)  # 首行缩进2字符≈18pt

    # 一级标题后加一个小空行
    if lv1:
        p_space = d.add_paragraph()
        set_para_fmt(p_space, space_before=0, space_after=0)
        run_sp = p_space.add_run('')
        set_run_fmt(run_sp, size_pt=5)


doc.save(DEST)
print(f"✅ 输出：{DEST}")
